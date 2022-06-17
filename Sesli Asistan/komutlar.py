from gtts import gTTS
from playsound import playsound
from random import choice
from lxml import html
from datetime import datetime
from googletrans import Translator
import os
import sys
import requests
import webbrowser
import speech_recognition as sr
import docx
import locale
import subprocess
import pyautogui
import time
import feedparser

locale.setlocale(locale.LC_ALL, 'tr_TR')

class Komut():
    def __init__(self,alinanses):
        self.ses = alinanses.upper()  #alınan sözcükleri büyük harfe dönüştürür.
        self.sesBloklari = self.ses.split() #alınan sözcüklerin tutuldugu liste

    def sesKaydi(self):
        r = sr.Recognizer()
        with sr.Microphone() as source:
            r.adjust_for_ambient_noise(source, duration=1)   #gürültüyü azaltır

            audio = r.listen(source, phrase_time_limit=4)    #dinleme süresi
            veri = ""
            try:
                veri = r.recognize_google(audio, language='tr-TR')
                print(veri)
                return veri
            except sr.UnknownValueError:
                self.konusma("Anlamadım")
            except sr.RequestError:
                self.konusma("Başarısız Oldu")

    def konusma(self, kelime):
        try:
            # Programın söyleyeceği sesler kaydedilir, oynatılır ve silinir.
            print(kelime)
            tts = gTTS(text=kelime, lang='tr')
            tts.save("ses.mp3")
            playsound("ses.mp3")
            os.remove("ses.mp3")
        except:
            self.konusma("Hata Alındı")

    def kapat(self):
        self.konusma("Kendine İyi bak")
        sys.exit()

    def hava_durumu(self):
        parse = feedparser.parse(
            "http://rss.accuweather.com/rss/liveweather_rss.asp?metric=1&locCode=EUR|TR|06420|ZONGULDAK|")
        parse = parse["entries"][0]["summary"]
        parse = parse.split()
        print(parse[2], parse[4], parse[5])
        # yazi = "Bugün hava {} derece ve {} gözüküyor.".format(parse[4].text, parse[2].text)
        # self.seslendirme(yazi)
        return (self.hava_durumu)

    def saat(self):
        try:
            yazi = datetime.now().strftime('%H:%M:%S')
            self.konusma(yazi)
        except:
            self.konusma("Hata Alındı")

    def word_olusturma(self):
        # Yeni bir Word dosyası açar.
        try:
            doc = docx.Document()
            self.konusma("Dosyanının ismini söyleyin")
            isim = self.sesKaydi()
            doc.save("{}.docx".format(isim))
            self.konusma("Tamamlandı")
        except:
            self.konusma("Bir hata oluştu")

    def excel_olusturma(self):
        # Yeni bir Excel dosyası açar.
        try:
            doc = docx.Document()
            self.konusma("Dosyanının ismini söyleyin")
            ad = self.sesKaydi()
            doc.save("{}.xlsx".format(ad))
            self.konusma("Tamamlandı")
        except:
            self.konusma("Hata Alındı")

    def powerpoint_olusturma(self):
        # Yeni bir PowerPoint dosyası açar.
        try:
            doc = docx.Document()
            self.konusma("Dosyanının ismini söyleyin")
            ad = self.sesKaydi()
            doc.save("{}.pptx".format(ad))
            self.konusma("Tamamlandı")
        except:
            self.konusma("Hata Alındı")

    def internetteArama(self):
        # Google'da istediğimiz bilgiyi arar.
        try:
            self.konusma("Senin için ne arayayım")
            bilgi = self.sesKaydi()
            url = "https://google.com/search?q="+bilgi
            webbrowser.get().open(url)
            self.konusma("{} için bulabildiklerim şunlar".format(bilgi))
        except:
            self.konusma("Hata Alındı")

    def internetteArama2(self):
        # Anlayamadığı durumlarda internetten aratır.
        try:
            url = "https://google.com/search?q="+self.ses
            webbrowser.get().open(url)
            self.konusma("{} için bulabildiklerim şunlar".format(self.ses))
        except:
            self.konusma("Hata Alındı")

    def tarih(self):
        # Tarih bilgisini söyler.
        try:
            an = datetime.now()
            yazi = "Tarih {} {} {} günlerden {}".format(an.day, an.strftime('%B'), an.year, an.strftime('%A'))
            self.konusma(yazi)
        except:
            self.konusma("Hata Alındı")

    def sohbet(self):
        # Sohbet mesajları.
        try:
            sozler = ["İyiyim sen nasılsın?", "Bugün çok iyiyim"," Daha özgür hissediyorum senden naber ?"]
            secenek = choice(sozler)
            self.konusma(secenek)
        except:
            self.konusma("Hata Alındı")

    def yaziYazma(self):
        # İmlecin bulunduğu yerde kullanıcıdan alınan kelimeleri yazar.
        try:
            self.konusma("Ne yazmamı istersin")
            yazi = self.sesKaydi()
            pyautogui.typewrite(yazi, interval=0.3)
        except:
            self.konusma("Hata Alındı")

    def ceviri_yap(self):
        # Türkçeden ingilizceye çeviri yapar
        try:
            self.konusma("İstenilen metni söyleyiniz")
            metin = self.sesKaydi()
            translator = Translator()
            text = translator.translate(metin, src="tr", dest="en").text
            self.konusma(text)

        except:
            self.konusma("Hata Alındı")

    #def ingilizce_cevir(self):
        # Türkçeden ingilizceye çeviri yapar
            # try:
            #self.konusma("İstenilen metni söyleyiniz")
            #metin = self.sesKaydi()
            #translator = Translator()
            #text = translator.translate(metin, src="en", dest="tr").text
            #self.konusma(text)

            #except:
            #self.konusma("Hata Alındı")

    def youtube(self):
        # İstenilen sonucu youtubeda arar.
        try:
            self.konusma("Ne aramak istiyorsun")
            bilgi = self.sesKaydi()
            url = "https://www.youtube.com/results?search_query=" + bilgi
            webbrowser.get().open(url)
            time.sleep(2)
            pyautogui.press("tab")
            pyautogui.press("enter")
            self.konusma("{} açıldı".format(bilgi))
        except:
            self.konusma("Hata Alındı")

    def harita(self):
        # İstenilen konumu haritada gösterir
        self.konusma("Konumu söyleyiniz")
        try:
            bilgi = self.sesKaydi()
            url = "https://www.google.com/maps/place/" + bilgi
            webbrowser.get().open(url)
            pyautogui.press("tab")
            pyautogui.press("enter")
            self.konusma("{} bulundu".format(bilgi))
        except:
            self.konusma("Hata Alındı")

    def haberler(self):
        # Güncel haberlerden 3 tanesini kullanıcıya okur
        try:
            url = ('https://www.haberturk.com/rss')
            haberler = feedparser.parse(url)
            for x in haberler.entries:
                self.konusma(x.title)
        except:
            self.konusma("Hata Alındı")

    def para(self):
        # Para birimleri arasında dönüşüm yapar ÖRNEK: "10 TL kaç Dolar"
        try:
            miktar = self.sesBloklari[0]
            istenen_para_birimi = self.sesBloklari[-1]
            if istenen_para_birimi == "TL" and self.sesBloklari[1] == "DOLAR":
                r = requests.get("https://tr.coinmill.com/TRY_USD.html?USD=" + miktar)
                tree = html.fromstring(r.content)
                sonuc = tree.xpath('//*[@id="currencyBox1"]/input')
                self.konusma("{} Dolar {} TL: ".format(miktar, str(sonuc[0].value)))
            elif istenen_para_birimi == "TL" and self.sesBloklari[1] == "KAÇ":
                r = requests.get("https://tr.coinmill.com/TRY_EUR.html?EUR=" + miktar[1])
                tree = html.fromstring(r.content)
                sonuc = tree.xpath('//*[@id="currencyBox1"]/input')
                self.konusma("{} {} TL: ".format(miktar, str(sonuc[0].value)))
            elif istenen_para_birimi == "DOLAR" and self.sesBloklari[1] == "TL":
                r = requests.get("https://tr.coinmill.com/USD_TRY.html?TRY=" + miktar)
                tree = html.fromstring(r.content)
                sonuc = tree.xpath('//*[@id="currencyBox1"]/input')
                self.konusma("{} TL {} dolar: ".format(miktar, str(sonuc[0].value)))
            elif istenen_para_birimi == "DOLAR" and self.sesBloklari[1] == "KAÇ":
                r = requests.get("https://tr.coinmill.com/EUR_USD.html?EUR=" + miktar[1])
                tree = html.fromstring(r.content)
                sonuc = tree.xpath('//*[@id="currencyBox1"]/input')
                self.konusma("{} {} dolar: ".format(miktar, str(sonuc[0].value)))
            elif istenen_para_birimi == "EURO" and self.sesBloklari[1] == "TL":
                r = requests.get("https://tr.coinmill.com/EUR_TRY.html?TRY=" + miktar)
                tree = html.fromstring(r.content)
                sonuc = tree.xpath('//*[@id="currencyBox1"]/input')
                self.konusma("{} TL {} EURO: ".format(miktar, str(sonuc[0].value)))
            elif istenen_para_birimi == "EURO" and self.sesBloklari[1] == "DOLAR":
                r = requests.get("https://tr.coinmill.com/USD_EUR.html?USD=" + miktar)
                tree = html.fromstring(r.content)
                sonuc = tree.xpath('//*[@id="currencyBox1"]/input')
                self.konusma("{} Dolar {} Euro: ".format(miktar, str(sonuc[0].value)))
        except:
            self.konusma("Hata Alındı")

    def komutBul(self):
        if "KENDINI KAPAT" in self.ses:
            self.kapat()
        elif "BILGISAYARI KAPAT" in self.ses:
            try:
                self.konusma("Şifreniz nedir")
                key = ""
                key = self.sesKaydi()
                if key == "SIFRE":
                    subprocess.Popen("C:\\Windows\\System32\\{}.exe".format("shutdown"))
                else:
                    self.konusma("Şifre Yanlış ")
            except:
                self.konusma("Hata Alındı")
        elif "BEKLE" in self.ses:
            self.konusma("Tamam")
            time.sleep(30)
        # Hava durumu
        elif "HAVA" in self.ses:
                self.hava_durumu()
        # Saati söyler
        elif "SAAT" in self.ses:
            self.saat()
        elif "TARIH" in self.ses:
            self.tarih()
        # Sohbet
        elif "NASILSIN" in self.ses or "NABER" in self.ses:
            self.sohbet()

        elif "TEŞEKKÜR" in self.ses or "SAĞ OL" in self.ses:
            self.konusma("Rica ederim")
        elif "SELAM" in self.ses or "MERHABA" in self.ses or "HI" in self.ses:
            self.konusma("Merhaba")
        elif "NAPIYORSUN" in self.ses:
            self.konusma("SENİ DINLIYORUM")
        elif "NERELISIN" in self.ses :
            self.konusma("FIKRIM YOK")
        elif "ADIN NE" in self.ses or "ISMIN NE" in self.ses:
            self.konusma("Adım CORTANA senin")
            self.sesKaydi()
            self.konusma("Ismin iyiymiş")
        elif "ISMININ ANLAMI" in self.ses:
            self.konusma("HALO oyunlarındaki yapay zekadan esinlenildi.")
        elif "CINSIYETIN NE" in self.ses or "INSAN MISIN" in self.ses or "SEN NESIN" in self.ses or "MAKINE MISIN" in self.ses:
            self.konusma("Ben bir asistanım.")
        elif "KAÇ YAŞINDASIN" in self.ses:
            self.konusma("FIKRIM YOK")
        elif "ŞARKI SÖYLE" in self.ses:
            self.konusma("lalalalaa")
        elif "INTERNET" in self.ses in self.ses:
            self.internetteArama()
        elif "EXCEL" in self.ses:
            self.excel_olusturma()
        elif "WORD" in self.ses:
            self.word_olusturma()
        elif "POWERPOINT" in self.ses:
            self.powerpoint_olusturma()
        elif "EKRAN GÖRÜNTÜSÜ" in self.ses:
            resim = pyautogui.screenshot("ekrangoruntusu.png")
            self.konusma("Ekran Görüntüsü alındı")
        elif "YAZI YAZ" in self.ses:
            self.yaziYazma()
        elif "HESAP MAKINESI" in self.ses:
            subprocess.Popen("C:\\Windows\\System32\\{}.exe".format("calc"))
        elif "SISTEM BILGI" in self.ses :
            subprocess.Popen("C:\\Windows\\System32\\{}.exe".format("dxdiag"))
        elif "RESIM" in self.ses:
            subprocess.Popen("C:\\Windows\\System32\\{}.exe".format("mspaint"))
        elif "BELGE" in self.ses:
            subprocess.Popen("C:\\Windows\\System32\\{}.exe".format("write"))
        elif "NOT" in self.ses:
            subprocess.Popen("C:\\Windows\\System32\\{}.exe".format("notepad"))
        elif "DENETIM MASASI" in self.ses:
            subprocess.Popen("C:\\Windows\\System32\\{}.exe".format("control"))
        elif "YOUTUBE" in self.ses:
            self.youtube()
        elif "HABER" in self.ses:
            self.haberler()
        elif "HARITA" in self.ses or "MAPS" in self.ses:
            self.harita()
        elif "TL" in self.ses or "DOLAR" in self.ses or "EURO" in self.ses:
            self.para()
        elif "ÇEVIRI" in self.ses:
            self.ceviri_yap()
        #elif "INGILIZCE ÇEVIRI" in self.ses:
         #   self.ingilizce_cevir()
        else:
            try:
                self.konusma("Anlamadım internette aramak ister misin ?")
                cevap = ""
                cevap = self.sesKaydi()
                if "evet" in cevap  \
                        or "ara" in cevap:
                    self.internetteArama2()
                else:
                    self.konusma("peki")
            except:
                self.konusma("Hata Alındı")